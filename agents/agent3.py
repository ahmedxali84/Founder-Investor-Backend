import os
import json
import re
import httpx
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from agents.llm import ask_llm_json, GROQ_API_KEY_AGENT3

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)   # deep professional blue
TEXT_GRAY    = RGBColor(0x1a, 0x1a, 0x2e)    # near-black, readable on white paper
LINK_COLOR   = RGBColor(0x1F, 0x4E, 0x79)    # same blue for hyperlinks


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR
    run.font.name = 'Calibri'
    add_bottom_border(p)
    return p


def bullet(doc, text):
    """Add a bullet point using a safe manually-formatted paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"• {text}")
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT_GRAY
    run.font.name = 'Calibri'
    return p


def add_hyperlink(paragraph, text, url):
    """
    Add a clickable hyperlink run inside an existing paragraph.
    Works with python-docx by injecting the proper OOXML relationship.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Styling: blue, underline
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '1F4E79')
    rPr.append(color_el)

    u_el = OxmlElement('w:u')
    u_el.set(qn('w:val'), 'single')
    rPr.append(u_el)

    sz_el = OxmlElement('w:sz')
    sz_el.set(qn('w:val'), '21')   # 10.5pt = 21 half-points
    rPr.append(sz_el)

    font_el = OxmlElement('w:rFonts')
    font_el.set(qn('w:ascii'), 'Calibri')
    font_el.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(font_el)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def bullet_with_link(doc, label_text, link_text, link_url):
    """
    Add a bullet where label text is plain, then a clickable hyperlink follows.
    e.g. "• TechFlixx Platform  →  [github.com/user/repo]"
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)

    label_run = p.add_run(f"• {label_text}  →  ")
    label_run.font.size = Pt(10.5)
    label_run.font.color.rgb = TEXT_GRAY
    label_run.font.name = 'Calibri'

    add_hyperlink(p, link_text, link_url)
    return p


def fetch_profile_image_bytes(url: str) -> bytes | None:
    """Download profile picture bytes from URL (GitHub avatar / LinkedIn picture)."""
    if not url:
        return None
    try:
        with httpx.Client(timeout=6.0, follow_redirects=True) as client:
            resp = client.get(url)
            if resp.status_code == 200 and resp.headers.get("content-type", "").startswith("image"):
                return resp.content
    except Exception as e:
        print(f"[Agent3] Profile image fetch error: {e}")
    return None


# ---------------------------------------------------------------------------
# Main Agent 3 function
# ---------------------------------------------------------------------------

def run_agent3(profile_data: dict, output_dir: str = "static/resumes", unique_id: str = "") -> str:
    """
    Take Agent 2's profile data, call Groq to polish it,
    then generate a formatted .docx resume and save it.

    Contact strategy:
      - GitHub URL present  → show GitHub URL in the contact line
      - GitHub absent + contact/email present → show contact/email
      - Both absent → skip contact line (investor-only prompt — not needed in resume)
    """

    # ---- 1. Groq polish ----
    prompt = f"""
You are a professional resume writer. Convert this founder's raw skill data
into a polished, investor-ready profile. Use clean, confident language,
no jargon left unexplained. Return ONLY valid JSON with this structure,
and nothing else:

{{
  "summary": "2-3 sentence professional summary",
  "skills": ["list of cleaned up skills"],
  "experience": ["bullet points, achievement-focused"],
  "projects": [
    {{
      "title": "Project name",
      "description": "One concise achievement-focused description",
      "github_url": "exact github repo url if available in profile_data repos, else empty string"
    }}
  ],
  "specialization_statement": "one sentence on their core strength"
}}

Founder data:
{json.dumps(profile_data, indent=2)}
"""
    polished = ask_llm_json(
        prompt,
        system_prompt="You are a professional resume writer that replies in valid JSON.",
        api_key=GROQ_API_KEY_AGENT3,
    )

    # ---- 2. Resolve contact info ----
    github_insights = profile_data.get("github_insights") or {}
    github_username = github_insights.get("username", "")
    github_url_raw = profile_data.get("github", "") or ""

    # Build the canonical GitHub profile URL
    if github_username:
        github_profile_url = f"https://github.com/{github_username}"
        github_display = f"github.com/{github_username}"
    elif github_url_raw:
        github_profile_url = github_url_raw.rstrip("/")
        github_display = github_profile_url.replace("https://", "").replace("http://", "")
    else:
        github_profile_url = ""
        github_display = ""

    # Contact fallback (email from linkedin_verified or custom_details)
    linkedin_verified = profile_data.get("linkedin_verified") or {}
    contact_email = (
        linkedin_verified.get("email")
        or (profile_data.get("custom_details") or {}).get("email")
        or ""
    )

    # ---- 3. Resolve profile image ----
    avatar_url = github_insights.get("avatar_url", "")
    # Prefer LinkedIn picture if present
    if linkedin_verified.get("picture"):
        avatar_url = linkedin_verified["picture"]
    image_bytes = fetch_profile_image_bytes(avatar_url)

    # ---- 4. Build the Word document ----
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # ---- Header table: Name + Profile Picture side-by-side ----
    header_table = doc.add_table(rows=1, cols=2)
    header_table.style = 'Table Grid'
    # Remove all borders from the header table
    tbl = header_table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Left cell — Name + Specialization
    left_cell = header_table.cell(0, 0)
    left_cell.width = Inches(4.5)
    name_p = left_cell.paragraphs[0]
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(profile_data.get("name", "Full Name"))
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = ACCENT_COLOR
    name_run.font.name = 'Calibri'

    spec_p = left_cell.add_paragraph()
    spec_p.paragraph_format.space_after = Pt(6)
    spec_run = spec_p.add_run(profile_data.get("specialization", ""))
    spec_run.font.size = Pt(12)
    spec_run.font.italic = True
    spec_run.font.color.rgb = TEXT_GRAY
    add_bottom_border(spec_p)

    # ---- Contact line (GitHub URL preferred, email fallback) ----
    if github_profile_url:
        contact_p = left_cell.add_paragraph()
        contact_p.paragraph_format.space_after = Pt(6)
        gh_label = contact_p.add_run("GitHub: ")
        gh_label.font.size = Pt(10.5)
        gh_label.font.bold = True
        gh_label.font.color.rgb = TEXT_GRAY
        gh_label.font.name = 'Calibri'
        add_hyperlink(contact_p, github_display, github_profile_url)
    elif contact_email:
        contact_p = left_cell.add_paragraph()
        contact_p.paragraph_format.space_after = Pt(6)
        c_run = contact_p.add_run(f"Contact: {contact_email}")
        c_run.font.size = Pt(10.5)
        c_run.font.color.rgb = TEXT_GRAY
        c_run.font.name = 'Calibri'
    # (else: no contact shown — both missing)

    # Right cell — Profile image (if available)
    right_cell = header_table.cell(0, 1)
    right_cell.width = Inches(1.25)
    right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if image_bytes:
        import io
        img_stream = io.BytesIO(image_bytes)
        img_para = right_cell.paragraphs[0]
        img_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            run = img_para.add_run()
            run.add_picture(img_stream, width=Inches(1.0))
        except Exception as e:
            print(f"[Agent3] Could not embed profile image: {e}")

    # ---- Professional Summary ----
    section_heading(doc, "Professional Summary")
    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.space_after = Pt(4)
    s_run = summary_p.add_run(polished.get("summary", ""))
    s_run.font.size = Pt(10.5)
    s_run.font.color.rgb = TEXT_GRAY

    # ---- Skills ----
    section_heading(doc, "Skills")
    skills_p = doc.add_paragraph()
    skills_p.paragraph_format.space_after = Pt(4)
    skills_text = "   |   ".join(polished.get("skills", []))
    sk_run = skills_p.add_run(skills_text)
    sk_run.font.size = Pt(10.5)
    sk_run.font.color.rgb = TEXT_GRAY
    sk_run.font.bold = True

    # ---- Experience ----
    section_heading(doc, "Experience")
    for exp in polished.get("experience", []):
        bullet(doc, exp)

    # ---- Projects (with clickable GitHub repo links) ----
    section_heading(doc, "Projects")
    projects = polished.get("projects", [])

    # Also pull real repo_details from Agent 2 github_insights for fallback URLs
    repo_details_map = {}
    for repo in (github_insights.get("repo_details") or []):
        repo_name = (repo.get("name") or "").lower()
        if repo_name and repo.get("url"):
            repo_details_map[repo_name] = repo["url"]

    if projects:
        for proj in projects:
            if isinstance(proj, dict):
                title = proj.get("title", "")
                description = proj.get("description", "")
                repo_url = proj.get("github_url", "").strip()

                # Fallback: try to find repo URL from Agent 2 github insights
                if not repo_url and title:
                    for repo_name, repo_link in repo_details_map.items():
                        if repo_name in title.lower() or title.lower() in repo_name:
                            repo_url = repo_link
                            break

                if repo_url:
                    # Show project title + description + clickable repo link
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.left_indent = Inches(0.25)
                    title_run = p.add_run(f"• {title}: ")
                    title_run.font.size = Pt(10.5)
                    title_run.font.bold = True
                    title_run.font.color.rgb = ACCENT_COLOR
                    title_run.font.name = 'Calibri'

                    desc_run = p.add_run(f"{description}  →  ")
                    desc_run.font.size = Pt(10.5)
                    desc_run.font.color.rgb = TEXT_GRAY
                    desc_run.font.name = 'Calibri'

                    # Clickable repo link
                    display_link = repo_url.replace("https://github.com/", "github.com/").replace("http://", "")
                    add_hyperlink(p, f"[{display_link}]", repo_url)
                else:
                    # No link available — plain bullet
                    full_text = f"{title}: {description}" if title else description
                    bullet(doc, full_text)
            else:
                # Legacy string format fallback
                bullet(doc, str(proj))
    else:
        # Fallback: use raw repo_details from github_insights
        for repo in (github_insights.get("repo_details") or [])[:6]:
            if isinstance(repo, dict):
                name = repo.get("name", "")
                desc = repo.get("description", "Repository project")
                url = repo.get("url", "")
                lang = repo.get("language", "")
                lang_str = f" [{lang}]" if lang else ""
                if url:
                    bullet_with_link(doc, f"{name}{lang_str}: {desc}", f"[github.com/{github_username}/{name}]", url)
                else:
                    bullet(doc, f"{name}{lang_str}: {desc}")

    # ---- Core Strength ----
    section_heading(doc, "Core Strength")
    spec_p2 = doc.add_paragraph()
    spec_run2 = spec_p2.add_run(polished.get("specialization_statement", ""))
    spec_run2.font.size = Pt(10.5)
    spec_run2.font.italic = True
    spec_run2.font.color.rgb = ACCENT_COLOR

    # ---- Education (if present) ----
    if profile_data.get("education"):
        section_heading(doc, "Education")
        edu_p = doc.add_paragraph()
        edu_run = edu_p.add_run(profile_data["education"])
        edu_run.font.size = Pt(10.5)
        edu_run.font.color.rgb = TEXT_GRAY

    # ---- Save ----
    os.makedirs(output_dir, exist_ok=True)
    raw_name = (profile_data.get("name") or "resume").strip()
    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", raw_name).strip("_") or "resume"
    safe_suffix = re.sub(r"[^A-Za-z0-9_-]+", "", unique_id)[:8]
    filename = f"resume_{safe_name}_{safe_suffix}.docx" if safe_suffix else f"resume_{safe_name}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    return output_path
