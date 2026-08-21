"""
Regression coverage for agent3.py's resume generation — the deterministic,
non-LLM parts: contact-line fallback logic, filename sanitization, and that
the generated .docx actually contains the profile's real content. ask_llm_json
(the "polish" step) and the profile-image fetch are both mocked; this isn't
re-testing Groq, it's testing that agent3 correctly assembles a real document
from whatever Groq returns.
"""
import os
from docx import Document as ReadDocument

import agents.agent3 as agent3


def _fake_polish(prompt, system_prompt="", api_key=""):
    return {
        "summary": "A concise, achievement-focused summary.",
        "skills": ["Python", "React"],
        "experience": ["Built a real-time chat platform used by 500 users."],
        "projects": [{"title": "Techflix", "description": "Founder-investor matching platform.", "github_url": ""}],
        "specialization_statement": "Full-stack engineer with an AI focus.",
    }


def _profile(**overrides):
    base = {
        "name": "Ahmed Ali",
        "specialization": "Fullstack, AI, Data Analysis",
        "github": "https://github.com/ahmedxali84",
        "github_insights": {"username": "ahmedxali84", "avatar_url": "", "repo_details": []},
    }
    base.update(overrides)
    return base


def _all_text(doc):
    """
    doc.paragraphs only covers paragraphs directly in the document body — the
    name/specialization/contact line all live inside the header table's left
    cell (see run_agent3's header_table), so a text check that skips
    doc.tables would silently never see them at all.
    """
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def test_run_agent3_produces_a_real_docx_with_the_founders_content(monkeypatch, tmp_path):
    monkeypatch.setattr(agent3, "ask_llm_json", _fake_polish)
    monkeypatch.setattr(agent3, "fetch_profile_image_bytes", lambda url: None)

    output_path = agent3.run_agent3(_profile(), output_dir=str(tmp_path), unique_id="user-123")

    assert output_path.endswith(".docx")
    assert os.path.isfile(output_path)
    doc = ReadDocument(output_path)
    all_text = _all_text(doc)
    assert "Ahmed Ali" in all_text
    assert "concise, achievement-focused summary" in all_text
    assert "Full-stack engineer with an AI focus" in all_text


def test_run_agent3_filename_sanitizes_name_and_truncates_unique_id(monkeypatch, tmp_path):
    monkeypatch.setattr(agent3, "ask_llm_json", _fake_polish)
    monkeypatch.setattr(agent3, "fetch_profile_image_bytes", lambda url: None)

    output_path = agent3.run_agent3(
        _profile(name="Ahmed  Ali! (CEO)"),
        output_dir=str(tmp_path),
        unique_id="12345678-abcd-ef00",
    )

    filename = os.path.basename(output_path)
    assert filename.startswith("resume_Ahmed_Ali")
    assert "!" not in filename and "(" not in filename
    # unique_id suffix is capped at 8 chars — "12345678", not the dash/rest.
    assert "12345678" in filename
    assert "abcd" not in filename


def test_run_agent3_falls_back_to_email_contact_when_no_github(monkeypatch, tmp_path):
    monkeypatch.setattr(agent3, "ask_llm_json", _fake_polish)
    monkeypatch.setattr(agent3, "fetch_profile_image_bytes", lambda url: None)

    profile = _profile(github="", github_insights={}, linkedin_verified={"email": "ahmed@example.com"})
    output_path = agent3.run_agent3(profile, output_dir=str(tmp_path), unique_id="u1")

    doc = ReadDocument(output_path)
    all_text = _all_text(doc)
    assert "ahmed@example.com" in all_text
    assert "GitHub:" not in all_text


def test_run_agent3_skips_contact_line_when_nothing_available(monkeypatch, tmp_path):
    monkeypatch.setattr(agent3, "ask_llm_json", _fake_polish)
    monkeypatch.setattr(agent3, "fetch_profile_image_bytes", lambda url: None)

    profile = _profile(github="", github_insights={})
    output_path = agent3.run_agent3(profile, output_dir=str(tmp_path), unique_id="u1")

    doc = ReadDocument(output_path)  # must not raise — the document is still valid
    all_text = _all_text(doc)
    assert "GitHub:" not in all_text
    assert "Contact:" not in all_text
